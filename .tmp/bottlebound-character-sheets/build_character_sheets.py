from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt, RGBColor


ROOT = Path("/Users/swb/projects/personal/bottlebound")
RULES = ROOT / "bottlebound_rules_final.md"
PORTRAITS = ROOT / "artifacts/bottlebound-character-sheets/portraits"
OUTPUT = ROOT / "artifacts/bottlebound-character-sheets/BOTTLEBOUND-character-abilities.docx"

FONT_DISPLAY = "Palatino Linotype"
FONT_BODY = "Aptos"


def clean_md(value: str) -> str:
    value = value.strip()
    value = value.replace("\\[", "[").replace("\\]", "]")
    value = value.replace("\\-", "-")
    value = re.sub(r"\*\*(.*?)\*\*", r"\1", value)
    value = re.sub(r"\*(.*?)\*", r"\1", value)
    return value.strip()


def parse_roster() -> list[dict]:
    lines = RULES.read_text(encoding="utf-8").splitlines()
    characters: list[dict] = []
    team = None
    current = None
    ability = None
    in_cards = False

    for line in lines:
        if line == "## 15. Character Ability Cards":
            in_cards = True
            continue
        if not in_cards:
            continue
        if line.startswith("## 16."):
            break
        if line in ("### DROW", "### DUERGAR"):
            team = line[4:].title()
            current = None
            ability = None
            continue
        match = re.match(r"^#### (.+?) — (.+)$", line)
        if match and team:
            current = {
                "team": team,
                "name": match.group(1),
                "role": match.group(2),
                "abilities": [],
            }
            characters.append(current)
            ability = None
            continue
        match = re.match(
            r"^HP (\d+) • Initiative ([+−\-]?\d+) • Basic Attack: (.+)$", line
        )
        if match and current:
            current["hp"] = match.group(1)
            current["initiative"] = match.group(2).replace("−", "-")
            current["basic_attack"] = match.group(3)
            continue
        if line.startswith("##### \\[ \\] ") and current:
            ability = {"name": line[len("##### \\[ \\] ") :]}
            current["abilities"].append(ability)
            continue
        if ability and line.startswith("|") and not re.match(r"^\|[-: ]+\|", line):
            cells = [clean_md(cell) for cell in line.strip().strip("|").split("|")]
            if len(cells) >= 2 and cells[0] and cells[0] not in ("Type",):
                ability[cells[0]] = cells[1]
            elif len(cells) >= 2 and cells[0] == "Type":
                ability["Type"] = cells[1]

    if len(characters) != 12 or any(len(c["abilities"]) != 2 for c in characters):
        raise RuntimeError(f"Expected 12 characters with 2 abilities, got {characters!r}")
    return characters


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=100, bottom=80, end=100) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_border(cell, color: str, size: int = 10, style: str = "single") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        el = borders.find(qn(f"w:{edge}"))
        if el is None:
            el = OxmlElement(f"w:{edge}")
            borders.append(el)
        el.set(qn("w:val"), style)
        el.set(qn("w:sz"), str(size))
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)


def set_table_geometry(table, widths_dxa: list[int], total_dxa: int, indent_dxa: int = 0) -> None:
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total_dxa))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tc_w = cell._tc.get_or_add_tcPr().find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                cell._tc.get_or_add_tcPr().append(tc_w)
            tc_w.set(qn("w:w"), str(widths_dxa[idx]))
            tc_w.set(qn("w:type"), "dxa")


def set_repeat_table_row_false(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def set_run(run, size: float, color: str, bold=False, italic=False, font=FONT_BODY) -> None:
    run.font.name = font
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), font)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), font)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold
    run.italic = italic


def configure_paragraph(paragraph, before=0, after=0, line=1.0, keep=False) -> None:
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line
    fmt.keep_with_next = keep


QUICK_TEXT = {
    "Backstab": ("STANDARD • THROW • 2 PACES", "Hit bottles take 1 damage and cannot use a Powerful Ability next turn. Allies can be hit too."),
    "Vanish": ("POWERFUL • SELF", "Move up to 4 paces. Physical balls cannot affect you until your next turn begins; Ability Attacks still can."),
    "Shapeshift": ("STANDARD • SELF • USE AT 2–3 HP", "Your maximum HP becomes 4, then heal 1 HP. It ends if damage takes you below 3 HP or you are Downed."),
    "Nature’s Renewal": ("STANDARD • SELF / ALLY • 4 PACES", "Restore 1 HP. Cannot target a Downed character. No Line of Sight needed."),
    "Lay on Hands": ("STANDARD • SELF / ALLY • 3 PACES", "Choose: heal 1 HP, or revive a Downed ally at 1 HP. No Line of Sight needed."),
    "Divine Shield": ("REACTION • SELF / ALLY • 3 PACES", "Damage Block: reduce that character’s remaining damage by 1. The hit and attached effects still resolve."),
    "Frostbind": ("STANDARD • ENEMY • 6 PACES • LOS", "On the target’s next turn, all movement is limited to 1 pace. They still take actions normally."),
    "Misty Escape": ("REACTION • SELF", "Attack Avoidance: take no damage or attached effects from the attack, then immediately move up to 2 paces."),
    "Arcane Bolt": ("STANDARD • ENEMY • 6 PACES • LOS", "Deal 1 automatic damage. No throw. Reactions can still prevent or modify it."),
    "Mirror Veil": ("REACTION • SELF", "Attack Avoidance: take no damage or attached effects from the attack against you."),
    "Inspiring Words": ("STANDARD • SELF / ALLY • 4 PACES", "Restore 1 HP. Cannot target a Downed character. No Line of Sight needed."),
    "Battle Hymn": ("POWERFUL • ALL DROW • 4 PACES", "Every living Drow ally currently in range receives +1 Move for the rest of the Match. Downing removes the bonus; Revival does not restore it."),
    "Hunter’s Mark": ("STANDARD • ENEMY • 6 PACES • LOS", "The next damaging attack against the target deals +1 damage. Expires at your next initiative if unused."),
    "Deadeye": ("STANDARD • ENEMY • 8 PACES • LOS", "Deal 1 automatic damage. No throw. Reactions can still prevent or modify it."),
    "Stunning Strike": ("STANDARD • THROW • 2 PACES", "Hit bottles take 1 damage and cannot use a Powerful Ability next turn. Allies can be hit too."),
    "Deflecting Palm": ("REACTION • SELF • PHYSICAL HIT", "Attack Avoidance: take no damage or attached effects, then redirect that same ball toward the thrower."),
    "Hold the Line": ("POWERFUL • FIGHTER + DUERGAR • 2 PACES", "At activation, fixed recipients reduce the first attack’s remaining damage against each of them by 1 until the Fighter’s next turn. Downing removes protection; Revival does not restore it."),
    "Shield Wall": ("REACTION • SELF / ALLY • 2 PACES", "Damage Block: reduce that character’s remaining damage by 1. The hit and attached effects still resolve."),
    "Brutal Shove": ("STANDARD • THROW • 2 PACES", "Hit bottles take 1 damage and are pushed up to 2 paces directly away. Allies can be hit too."),
    "Rampage": ("POWERFUL • THROW • 2 PACES", "Move up to twice your current Move, then throw. Hit bottles take 1 damage and are pushed up to 2 paces directly away."),
    "Hex": ("STANDARD • ENEMY • 6 PACES • LOS", "The next damaging attack against the target deals +1 damage; then their movement is capped at 1 pace next turn. Expires at your next initiative if unused."),
    "Eldritch Blast": ("STANDARD • ENEMY • 6 PACES • LOS", "Deal 1 automatic damage. No throw. Reactions can still prevent or modify it."),
    "Blessing of Battle": ("STANDARD • ALLY • 4 PACES", "The ally receives +1 Move for the rest of the Match. Downing removes the bonus; Revival does not restore it. No Line of Sight needed."),
    "Revivify": ("STANDARD • DOWNED ALLY • 3 PACES", "Revive the ally at 1 HP. They return at their normal initiative; no immediate extra turn."),
}


def add_rule(paragraph, color: str, size: int = 8) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def portrait_path(character: dict) -> Path:
    slug = f"{character['team']}-{character['name']}".lower().replace(" ", "-")
    return PORTRAITS / f"{slug}.png"


def add_ability(cell, ability: dict, palette: dict) -> None:
    meta_text, cue_text = QUICK_TEXT[ability["name"]]
    if ability.get("Type") == "Standard":
        meta_text = f"{meta_text} • 1 ACTION"
    elif ability.get("Type") == "Powerful":
        meta_text = f"{meta_text} • 2 ACTIONS"
    else:
        meta_text = f"{meta_text} • NO ACTION"
    heading = cell.add_paragraph()
    configure_paragraph(heading, before=3.5, after=1.1, line=1.0, keep=True)
    name_run = heading.add_run(f"[ ] {ability['name'].upper()}")
    set_run(name_run, 10.4, palette["accent"], bold=True, font=FONT_DISPLAY)
    meta = cell.add_paragraph()
    configure_paragraph(meta, after=1.2, line=1.0, keep=True)
    set_run(meta.add_run(meta_text), 7.9, palette["muted"], bold=True)

    effect = cell.add_paragraph()
    configure_paragraph(effect, after=1.6, line=1.03)
    set_run(effect.add_run(cue_text), 9.0, palette["ink"])


def add_character_card(cell, character: dict, palette: dict) -> None:
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    set_cell_shading(cell, palette["paper"])
    set_cell_border(cell, palette["border"], size=12, style="double")
    set_cell_margins(cell, top=90, start=110, bottom=70, end=110)

    # Remove the empty starter paragraph only after the nested table exists.
    top = cell.add_table(rows=1, cols=2)
    top.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_geometry(top, [1700, 3500], 5200, 0)
    for top_cell in top.rows[0].cells:
        set_cell_shading(top_cell, palette["paper"])
        set_cell_margins(top_cell, top=0, start=0, bottom=0, end=80)
        set_cell_border(top_cell, palette["paper"], size=0, style="nil")

    image_p = top.cell(0, 0).paragraphs[0]
    image_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    configure_paragraph(image_p, after=0)
    picture = image_p.add_run().add_picture(
        str(portrait_path(character)), width=Mm(27.5), height=Mm(34.0)
    )
    alt = f"{character['team']} {character['name']}, {character['role']} character portrait"
    picture._inline.docPr.set("descr", alt)
    picture._inline.docPr.set("title", alt)

    info = top.cell(0, 1)
    p = info.paragraphs[0]
    configure_paragraph(p, after=0.5, line=0.95, keep=True)
    set_run(p.add_run(character["name"].upper()), 14.0, palette["title"], bold=True, font=FONT_DISPLAY)
    p = info.add_paragraph()
    configure_paragraph(p, after=1.0, line=0.9, keep=True)
    set_run(p.add_run(character["role"].upper()), 8.0, palette["accent"], bold=True)
    p = info.add_paragraph()
    configure_paragraph(p, after=0.7, line=0.9, keep=True)
    stats = f"HP {character['hp']}   ◆   INIT {character['initiative']}"
    set_run(p.add_run(stats), 9.0, palette["ink"], bold=True)
    p = info.add_paragraph()
    configure_paragraph(p, after=0, line=0.9)
    set_run(p.add_run(f"BASIC  {character['basic_attack']}"), 7.8, palette["muted"], bold=True)

    # Drop the cell's original empty paragraph.
    first = cell.paragraphs[0]
    first._element.getparent().remove(first._element)

    divider = cell.add_paragraph()
    configure_paragraph(divider, before=0.2, after=0.2)
    add_rule(divider, palette["border"], 7)

    for ability in character["abilities"]:
        add_ability(cell, ability, palette)


PALETTES = {
    "Drow": {
        "paper": "F2EDF4",
        "ink": "1D1723",
        "title": "21152B",
        "accent": "673D83",
        "muted": "5E5364",
        "border": "8F72A1",
        "banner": "25152F",
        "banner_text": "F6EAFB",
        "kicker": "CDB1DF",
    },
    "Duergar": {
        "paper": "F1ECE2",
        "ink": "211C18",
        "title": "2B201B",
        "accent": "91401F",
        "muted": "62564E",
        "border": "A66B43",
        "banner": "30251F",
        "banner_text": "FFF2DC",
        "kicker": "D8A06B",
    },
}


def set_page_banner(doc: Document, team: str, palette: dict) -> None:
    banner = doc.add_table(rows=1, cols=2)
    banner.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_geometry(banner, [3150, 12750], 15900, 0)
    for c in banner.rows[0].cells:
        set_cell_shading(c, palette["banner"])
        set_cell_margins(c, top=85, start=130, bottom=85, end=130)
        set_cell_border(c, palette["banner"], 0, "nil")
        c.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    left = banner.cell(0, 0).paragraphs[0]
    configure_paragraph(left, after=0, line=0.9)
    set_run(left.add_run("BOTTLEBOUND"), 12.5, palette["banner_text"], bold=True, font=FONT_DISPLAY)
    right = banner.cell(0, 1).paragraphs[0]
    right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    configure_paragraph(right, after=0, line=0.9)
    set_run(right.add_run(f"{team.upper()}  ◆  PLAYER QUICK SHEET"), 11.5, palette["kicker"], bold=True, font=FONT_DISPLAY)


def add_footer(section, palette: dict) -> None:
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    configure_paragraph(p, after=0, line=1.0)
    set_run(
        p.add_run("PLAYER QUICK CUES • One use per ability • Referee resolves timing, edge cases, and full rules"),
        7.3,
        palette["muted"],
        italic=True,
    )


def build() -> None:
    characters = parse_roster()
    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Mm(297)
    section.page_height = Mm(210)
    section.top_margin = Mm(6)
    section.bottom_margin = Mm(6)
    section.left_margin = Mm(6)
    section.right_margin = Mm(6)
    section.header_distance = Mm(2)
    section.footer_distance = Mm(2)
    add_footer(section, PALETTES["Drow"])

    normal = doc.styles["Normal"]
    normal.font.name = FONT_BODY
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT_BODY)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_BODY)
    normal.font.size = Pt(6.2)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.line_spacing = 1.0

    for page_index, team in enumerate(("Drow", "Duergar")):
        if page_index:
            doc.add_page_break()
        palette = PALETTES[team]
        set_page_banner(doc, team, palette)
        spacer = doc.add_paragraph()
        configure_paragraph(spacer, after=0.8)

        grid = doc.add_table(rows=2, cols=3)
        grid.alignment = WD_TABLE_ALIGNMENT.CENTER
        set_table_geometry(grid, [5300, 5300, 5300], 15900, 0)
        for row in grid.rows:
            set_repeat_table_row_false(row)
        team_chars = [c for c in characters if c["team"] == team]
        for idx, character in enumerate(team_chars):
            add_character_card(grid.cell(idx // 3, idx % 3), character, palette)

    props = doc.core_properties
    props.title = "BOTTLEBOUND Character Abilities"
    props.subject = "Print-ready team ability sheets"
    props.author = "BOTTLEBOUND"
    props.keywords = "BOTTLEBOUND, character abilities, Drow, Duergar, printable"
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
